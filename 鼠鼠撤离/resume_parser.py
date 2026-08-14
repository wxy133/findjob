import json
import os
import re


class ResumeParser:
    """Lightweight local parser with structured education records."""

    DEGREE_ALIASES = (
        ("博士", ("博士", "博士生")),
        ("硕士", ("硕士", "研究生")),
        ("本科", ("本科", "学士")),
        ("大专", ("大专", "专科")),
        ("MBA", ("MBA",)),
        ("高中", ("高中",)),
        ("中专", ("中专",)),
    )
    SCHOOL_PATTERN = re.compile(
        r"(?<![\u4e00-\u9fffA-Za-z0-9])"
        r"([\u4e00-\u9fffA-Za-z0-9·]{2,32}?(?:大学|学院|学校|研究院))"
    )

    def __init__(self):
        self.text = ""
        self._reset_results()

    def _reset_results(self):
        self.name = self.phone = self.email = self.target_position = ""
        self.skills = []
        self.experience = []
        self.education = []
        self.keywords = []

    def load_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            self._load_pdf(file_path)
        elif ext in (".docx", ".doc"):
            self._load_docx(file_path)
        elif ext == ".txt":
            self._load_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
        self._parse()
        return self.get_summary()

    def _load_pdf(self, file_path):
        try:
            import pdfplumber
        except ImportError as exc:
            raise ImportError("请先安装 pdfplumber: pip install pdfplumber") from exc
        with pdfplumber.open(file_path) as pdf:
            self.text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    def _load_docx(self, file_path):
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("请先安装 python-docx: pip install python-docx") from exc
        doc = Document(file_path)
        lines = [paragraph.text for paragraph in doc.paragraphs]
        lines.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        self.text = "\n".join(lines)

    def _load_txt(self, file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            self.text = file.read()

    def _parse(self):
        self._reset_results()
        self._extract_name()
        self._extract_phone()
        self._extract_email()
        self._extract_skills()
        self._extract_experience()
        self._extract_education()
        self._extract_target_position()
        self._extract_keywords()

    def _extract_name(self):
        for line in self.text.splitlines()[:10]:
            candidate = line.strip()
            if re.fullmatch(r"[\u4e00-\u9fff]{2,4}", candidate):
                self.name = candidate
                return
        match = re.search(r"姓名\s*[:：\s]*([\u4e00-\u9fff]{2,4})", self.text)
        if match:
            self.name = match.group(1)

    def _extract_phone(self):
        match = re.search(r"1[3-9]\d{9}", self.text)
        self.phone = match.group() if match else ""

    def _extract_email(self):
        match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", self.text)
        self.email = match.group() if match else ""

    def _extract_skills(self):
        skills = (
            "Python", "Java", "C++", "C#", "JavaScript", "TypeScript", "Go", "Rust", "PHP",
            "React", "Vue", "Angular", "Node.js", "Django", "Flask", "Spring", "SpringBoot",
            "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle", "SQL", "Linux", "Docker",
            "Kubernetes", "AWS", "阿里云", "腾讯云", "Git", "SVN", "Jenkins", "CI/CD",
            "TensorFlow", "PyTorch", "机器学习", "深度学习", "数据分析", "HTML", "CSS",
            "微服务", "分布式", "高并发", "性能优化", "产品经理", "UI设计", "运营",
        )
        text_lower = self.text.lower()
        self.skills = list(dict.fromkeys(skill for skill in skills if skill.lower() in text_lower))

    def _extract_experience(self):
        pattern = re.compile(
            r"([\u4e00-\u9fffA-Za-z0-9]{2,30}?(?:公司|科技|集团|股份|有限|工作室|研究院))"
            r"\s*[\s\S]{0,50}?(\d{4}[年/-]\d{1,2}[月/-]?)\s*[至到-]\s*"
            r"(\d{4}[年/-]\d{1,2}[月/-]?|至今|现在)"
        )
        self.experience = [
            {"company": company, "period": f"{start} - {end}"}
            for company, start, end in pattern.findall(self.text)
        ]

    def _extract_education(self):
        """Pair each detected degree to its closest school, rather than flattening."""
        degrees = []
        for canonical, aliases in self.DEGREE_ALIASES:
            for alias in aliases:
                degrees.extend((match.start(), canonical) for match in re.finditer(re.escape(alias), self.text, re.I))
        degrees.sort(key=lambda item: item[0])
        schools = [(match.start(1), match.group(1).strip()) for match in self.SCHOOL_PATTERN.finditer(self.text)]
        used = set()
        records = []
        for degree_position, degree in degrees:
            candidates = [item for item in schools if item[0] not in used]
            nearby = [item for item in candidates if abs(item[0] - degree_position) <= 240]
            pool = nearby or candidates
            school = ""
            if pool:
                school_position, candidate = min(pool, key=lambda item: abs(item[0] - degree_position))
                if abs(school_position - degree_position) <= 320:
                    school = candidate
                    used.add(school_position)
            records.append({"degree": degree, "school": school})
        for school_position, school in schools:
            if school_position not in used:
                records.append({"degree": "", "school": school})
        self.education = records

    def _extract_target_position(self):
        for pattern in (
            r"求职意向[:：\s]*([\u4e00-\u9fffA-Za-z0-9/、\s]{2,30})",
            r"目标岗位[:：\s]*([\u4e00-\u9fffA-Za-z0-9/、\s]{2,30})",
            r"应聘职位[:：\s]*([\u4e00-\u9fffA-Za-z0-9/、\s]{2,30})",
            r"期望职位[:：\s]*([\u4e00-\u9fffA-Za-z0-9/、\s]{2,30})",
        ):
            match = re.search(pattern, self.text)
            if match:
                self.target_position = match.group(1).strip()
                return

    def _extract_keywords(self):
        self.keywords = self.skills.copy()
        if self.target_position:
            self.keywords.append(self.target_position)
        self.keywords.extend(item["degree"] for item in self.education if item["degree"])
        self.keywords = list(dict.fromkeys(self.keywords))

    def get_summary(self):
        return {
            "name": self.name, "phone": self.phone, "email": self.email,
            "skills": self.skills, "experience": self.experience,
            "education": self.education, "target_position": self.target_position,
            "keywords": self.keywords, "raw_text": self.text[:2000],
        }

    def get_search_keywords(self):
        return self.target_position or " ".join(self.skills[:3])

    def to_json(self):
        return json.dumps(self.get_summary(), ensure_ascii=False, indent=2)
