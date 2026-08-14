import os
import re
import json


class ResumeParser:
    def __init__(self):
        self.text = ""
        self.name = ""
        self.phone = ""
        self.email = ""
        self.skills = []
        self.experience = []
        self.education = []
        self.target_position = ""
        self.keywords = []

    def load_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            self._load_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            self._load_docx(file_path)
        elif ext == '.txt':
            self._load_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
        self._parse()
        return self.get_summary()

    def _load_pdf(self, file_path):
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            self.text = text
        except ImportError:
            raise ImportError("请先安装 pdfplumber: pip install pdfplumber")

    def _load_docx(self, file_path):
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            self.text = text
        except ImportError:
            raise ImportError("请先安装 python-docx: pip install python-docx")

    def _load_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            self.text = f.read()

    def _parse(self):
        self._extract_name()
        self._extract_phone()
        self._extract_email()
        self._extract_skills()
        self._extract_experience()
        self._extract_education()
        self._extract_target_position()
        self._extract_keywords()

    def _extract_name(self):
        lines = self.text.split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) >= 2 and len(line) <= 4:
                if re.match(r'^[\u4e00-\u9fa5]{2,4}$', line):
                    self.name = line
                    return
        match = re.search(r'姓\s*名[:：\s]*([\u4e00-\u9fa5]{2,4})', self.text)
        if match:
            self.name = match.group(1)

    def _extract_phone(self):
        match = re.search(r'1[3-9]\d{9}', self.text)
        if match:
            self.phone = match.group()

    def _extract_email(self):
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', self.text)
        if match:
            self.email = match.group()

    def _extract_skills(self):
        skill_keywords = [
            'Python', 'Java', 'C++', 'C#', 'JavaScript', 'TypeScript', 'Go', 'Rust', 'PHP',
            'React', 'Vue', 'Angular', 'Node.js', 'Django', 'Flask', 'Spring', 'SpringBoot',
            'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL',
            'Linux', 'Docker', 'Kubernetes', 'AWS', '阿里云', '腾讯云',
            'Git', 'SVN', 'Jenkins', 'CI/CD',
            'TensorFlow', 'PyTorch', '机器学习', '深度学习', '数据分析',
            'HTML', 'CSS', 'Sass', 'Less', 'Webpack', 'Vite',
            '微服务', '分布式', '高并发', '性能优化',
            '产品经理', 'UI设计', '运营', '销售', '财务', '人事', '行政'
        ]
        found = []
        text_lower = self.text.lower()
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                found.append(skill)
        self.skills = list(set(found))

    def _extract_experience(self):
        patterns = [
            r'(\d{4}[年/-]\d{1,2}[月/-]?)\s*[至到-]\s*(\d{4}[年/-]\d{1,2}[月/-]?|至今|现在)\s*[\s\S]{0,100}?([\u4e00-\u9fa5A-Za-z0-9]{2,30}?(?:公司|科技|集团|股份|有限|工作室|研究院|学校|大学))',
            r'([\u4e00-\u9fa5A-Za-z0-9]{2,30}?(?:公司|科技|集团|股份|有限|工作室|研究院|学校|大学))\s*[\s\S]{0,50}?(\d{4}[年/-]\d{1,2}[月/-]?)\s*[至到-]\s*(\d{4}[年/-]\d{1,2}[月/-]?|至今|现在)',
        ]
        for pattern in patterns:
            matches = re.findall(pattern, self.text)
            if matches:
                for m in matches:
                    exp = {
                        'company': m[0] if '公司' in m[0] or '科技' in m[0] else m[2],
                        'period': f"{m[1] if not '公司' in m[0] and not '科技' in m[0] else m[0]} - {m[2] if not '公司' in m[0] and not '科技' in m[0] else m[1]}"
                    }
                    self.experience.append(exp)
                if self.experience:
                    break

    def _extract_education(self):
        edu_keywords = ['本科', '硕士', '博士', '大专', '高中', '中专', '学士', '研究生', 'MBA']
        for edu in edu_keywords:
            if edu in self.text:
                self.education.append(edu)
        univ_pattern = r'([\u4e00-\u9fa5A-Za-z0-9]{2,20}?(?:大学|学院|学校))'
        univs = re.findall(univ_pattern, self.text)
        if univs:
            self.education.extend(list(set(univs)))

    def _extract_target_position(self):
        patterns = [
            r'求职意向[:：\s]*([\u4e00-\u9fa5A-Za-z0-9/、\s]{2,30})',
            r'目标岗位[:：\s]*([\u4e00-\u9fa5A-Za-z0-9/、\s]{2,30})',
            r'应聘职位[:：\s]*([\u4e00-\u9fa5A-Za-z0-9/、\s]{2,30})',
            r'期望职位[:：\s]*([\u4e00-\u9fa5A-Za-z0-9/、\s]{2,30})',
        ]
        for pattern in patterns:
            match = re.search(pattern, self.text)
            if match:
                pos = match.group(1).strip()
                if 2 <= len(pos) <= 30:
                    self.target_position = pos
                    return

    def _extract_keywords(self):
        self.keywords = self.skills.copy()
        if self.target_position:
            self.keywords.append(self.target_position)
        for edu in self.education:
            if edu in ['本科', '硕士', '博士', '大专']:
                self.keywords.append(edu)
        self.keywords = list(set(self.keywords))

    def get_summary(self):
        return {
            'name': self.name,
            'phone': self.phone,
            'email': self.email,
            'skills': self.skills,
            'experience': self.experience,
            'education': self.education,
            'target_position': self.target_position,
            'keywords': self.keywords,
            'raw_text': self.text[:2000] if len(self.text) > 2000 else self.text
        }

    def get_search_keywords(self):
        if self.target_position:
            return self.target_position
        if self.skills:
            return ' '.join(self.skills[:3])
        return ''

    def to_json(self):
        return json.dumps(self.get_summary(), ensure_ascii=False, indent=2)


# Keep each education record paired as {degree, school}. The original parser
# flattened all degrees and schools into one list, which lost their relation.
def _extract_education_paired(self):
    degree_aliases = (
        ("博士", ("博士", "博士生")), ("硕士", ("硕士", "研究生")),
        ("本科", ("本科", "学士")), ("大专", ("大专", "专科")),
        ("MBA", ("MBA",)), ("高中", ("高中",)), ("中专", ("中专",)),
    )
    degrees = []
    for canonical, aliases in degree_aliases:
        for alias in aliases:
            degrees.extend((match.start(), canonical) for match in re.finditer(re.escape(alias), self.text, re.I))
    degrees.sort(key=lambda item: item[0])
    school_pattern = re.compile(
        r"(?<![\u4e00-\u9fffA-Za-z0-9])"
        r"([\u4e00-\u9fffA-Za-z0-9·]{2,32}?(?:大学|学院|学校|研究院))"
    )
    schools = [(match.start(1), match.group(1).strip()) for match in school_pattern.finditer(self.text)]
    used = set()
    records = []
    for degree_position, degree in degrees:
        candidates = [item for item in schools if item[0] not in used]
        nearby = [item for item in candidates if abs(item[0] - degree_position) <= 240]
        pool = nearby or candidates
        school = ""
        if pool:
            school_position, candidate = min(pool, key=lambda item: abs(item[0] - degree_position))
            if abs(school_position - degree_position) <= 320:
                school = candidate
                used.add(school_position)
        records.append({"degree": degree, "school": school})
    for school_position, school in schools:
        if school_position not in used:
            records.append({"degree": "", "school": school})
    self.education = records


def _extract_keywords_paired(self):
    self.keywords = self.skills.copy()
    if self.target_position:
        self.keywords.append(self.target_position)
    self.keywords.extend(item["degree"] for item in self.education if item.get("degree"))
    self.keywords = list(dict.fromkeys(self.keywords))


ResumeParser._extract_education = _extract_education_paired
ResumeParser._extract_keywords = _extract_keywords_paired
